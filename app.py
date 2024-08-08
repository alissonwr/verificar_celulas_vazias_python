from flask import Flask, request, render_template, redirect, url_for, send_from_directory, flash
import openpyxl
from openpyxl.styles import PatternFill
from docx import Document
import os

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'outputs'
app.secret_key = 'supersecretkey'

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

def verificar_celulas_vazias(caminho_arquivo, nome_planilha):
    wb = openpyxl.load_workbook(caminho_arquivo)
    ws = wb[nome_planilha]

    intervalos_mesclados = ws.merged_cells.ranges
    celulas_vazias = []

    def esta_em_intervalo_mesclado(cell):
        for intervalo in intervalos_mesclados:
            if cell.coordinate in intervalo:
                return True
        return False

    for row in ws.iter_rows():
        for cell in row:
            if cell.value is None and not esta_em_intervalo_mesclado(cell):
                celulas_vazias.append(cell.coordinate)

    fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
    for celula in celulas_vazias:
        ws[celula].fill = fill

    output_excel = os.path.join(app.config['OUTPUT_FOLDER'], 'output_verificacao.xlsx')
    wb.save(output_excel)

    doc = Document()
    doc.add_heading('Relatório de Células Vazias', level=1)

    if celulas_vazias:
        doc.add_paragraph("Células vazias encontradas nas seguintes posições:")
        for celula in celulas_vazias:
            doc.add_paragraph(celula)
    else:
        doc.add_paragraph("Nenhuma célula vazia encontrada.")

    output_word = os.path.join(app.config['OUTPUT_FOLDER'], 'relatorio_celulas_vazias.docx')
    doc.save(output_word)

    return output_excel, output_word

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        file = request.files['file']
        if file:
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
            file.save(filepath)

            wb = openpyxl.load_workbook(filepath)
            sheet_names = wb.sheetnames

            return render_template('select_sheet.html', sheet_names=sheet_names, file_name=file.filename)

    return render_template('index.html')

@app.route('/analyze', methods=['POST'])
def analyze():
    file_name = request.form['file_name']
    sheet_name = request.form['sheet_name']

    if file_name and sheet_name:
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], file_name)

        try:
            output_excel, output_word = verificar_celulas_vazias(filepath, sheet_name)
            return render_template('result.html', excel_report=output_excel, word_report=output_word)
        except KeyError:
            flash(f'A planilha "{sheet_name}" não existe no arquivo.')
            return redirect(url_for('index'))

    return redirect(url_for('index'))

@app.route('/download/<filename>')
def download_file(filename):
    return send_from_directory(app.config['OUTPUT_FOLDER'], filename, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True, port=8000)
